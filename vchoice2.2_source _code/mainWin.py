
from random import randint

# from test import *
import sys
import pandas
from PyQt5.QtCore import Qt, QSize, QRect, QUrl, QFileInfo
from PyQt5 import QtGui
from PyQt5.QtWidgets import *
from PyQt5.QtWebEngineWidgets import QWebEngineView
from analyse import StyleBoxfileModel,YearBoxfileModel,TopFilmModel,TopActorModel
from sift import sift

from predict_boxfile import *
from docx import Document
from docx.shared import Inches
from PIL import Image
import shutil, os


class DeleteFiles(object):
    def __init__(self, pathDir):
        self.pathDir = pathDir

    def delete_files(self):
        os.chdir(self.pathDir)
        fileList = list(os.listdir())
        for file in fileList:
            if os.path.isfile(file):
                os.remove(file)
                print("delete successfully")
            else:
                shutil.rmtree(file)



class LeftTabWidget(QDialog):

    def __init__(self, id, date, authority = 1):
        super(LeftTabWidget, self).__init__()
        self.id = id
        self.date = date
        self.authority = authority
        self.graphList = []

        self.setWindowFlags(Qt.WindowMinimizeButtonHint|Qt.WindowCloseButtonHint)
        self.setWindowTitle("Vchoice")
        # 美化样式表
        self.setWindowIcon(QtGui.QIcon('./image/flash.ico'))
        self.Stylesheet = """
        QListWidget, QListView, QTreeWidget, QTreeView {
            outline: 0px;
        }
        /*标题栏*/
        TitleBar {
            background-color: rgb(54, 157, 180);
        }
/*最小化最大化关闭按钮通用默认背景*/
        #buttonMinimum,#buttonMaximum,#buttonClose {
            border: none;
            background-color: rgb(54, 157, 180);
        }
        QListWidget {
            min-width: 120px;
            max-width: 120px;
            color: white;
            background: rgb(61,61,61);
        }
        QListWidget::item:selected {
            background: rgb(52, 52, 52);
            border-left: 2px solid rgb(255, 200, 1);
        }

        HistoryPanel::item:hover {
            background: rgb(52, 52, 52);
        }


        QStackedWidget {
            background-image: url(:/new/prefix1/source/picture/极简灰白微立体通用PPT模板2.png);
        }
        QPushButton#loginButton {
            color: white; 
            background-color: black;
        }
        #loginButton:hover {
            background-color: rgb(255, 200, 1); 
        }
        #loginButton:pressed {
            background-color: #e57373; 
        }

        """


        self.setStyleSheet(self.Stylesheet)
        self.resize(1660, 843)

        layout =QVBoxLayout(self, spacing=0)
        childLayout = QHBoxLayout(self, spacing=0)

        childLayout.setContentsMargins(0, 0, 0, 0)
        self.listWidget = QListWidget(self)
        childLayout.addWidget(self.listWidget)

        self.stackedWidget = QStackedWidget(self)

        # self.stackedWidget.setStyleSheet("background: rgb(61, 61, 61);")
        childLayout.addWidget(self.stackedWidget)

        headLayout = QHBoxLayout(self, spacing=0)
        headLayout.setContentsMargins(0, 0, 0, 0)

        logal = QLabel()
        logal.setStyleSheet("image: url(:/new/prefix1/source/picture/稿定设计导出-20181211-154931.png); background-color: black")
        logal.setAlignment(Qt.AlignLeft)
        logal.setScaledContents(True)
        logal.setMinimumSize(130,40)
        logal.setMaximumSize(130,40)

        black = QLabel('')
        black.setStyleSheet("background-color: black;")

        if self.authority == 2:
            black.setText('游客模式')
            black.setStyleSheet('color : white; background-color: black;font: 14pt \"黑体\"')

        if self.authority == 2:
            loginButton = QPushButton("立即登录",objectName="loginButton")
        else:
            loginButton = QPushButton("注销", objectName="loginButton")
        loginButton.setMaximumSize(130,40)
        loginButton.clicked.connect(self.login)
        headLayout.addWidget(logal)
        headLayout.addWidget(black)
        headLayout.addWidget(loginButton)


        layout.setContentsMargins(0, 0, 0, 0)
        layout.addLayout(headLayout)
        layout.addLayout(childLayout)

        self.stack1 = QWidget()
        self.stack2 = QWidget()
        self.stack3 = QWidget()
        self.stack4 = QWidget()
        if self.authority != 2:
            self.stack5 = QWidget()
            self.stack6 = QWidget()
            self.stack7 = QWidget()
        self.seasonBox = QComboBox()
        self.seasonBox.setMaximumSize(150, 50)
        # self.seasonBox.setStyleSheet('color: rgb(199, 199 ,199)')
        self.seasonInfor = ['', '第一季度', '第二季度', '第三季度', '第四季度']
        self.seasonBox.addItems(self.seasonInfor)
        self.seasonBox.currentIndexChanged.connect(self.mouthchange)

        self.mouthBox = QComboBox()
        self.mouthBox.setMaximumSize(150, 50)
        # self.mouthBox.setStyleSheet('color: rgb(199, 199 ,199)')
        self.mouthBox.addItems(['','','',''])

        self.stack1UI()
        self.stack2UI()
        self.stack3UI()
        self.stack4UI()
        if self.authority != 2:
            self.stack5UI()
            self.stack6UI()
            self.stack7UI()

        self.stackedWidget.addWidget(self.stack1)
        self.stackedWidget.addWidget(self.stack2)
        self.stackedWidget.addWidget(self.stack3)
        self.stackedWidget.addWidget(self.stack4)
        if self.authority != 2:
            self.stackedWidget.addWidget(self.stack5)
            self.stackedWidget.addWidget(self.stack6)
            self.stackedWidget.addWidget(self.stack7)

        self.initUi()

    def login(self):
        self.accept()
        self.close()
        return True




    def closeEvent(self, event):
        if self.authority == 1:
            reply = QMessageBox.question(self, '警告', '退出后图表将不保存,\n你确认要退出吗？',
                                                   QMessageBox.Yes, QMessageBox.No)
            if reply == QMessageBox.Yes:
                path = "Users"
                delete = DeleteFiles(path)
                delete.delete_files()
                event.accept()
            else:
                event.ignore()
        self.reject

    def mouthchange(self):
        season = self.seasonBox.currentText()
        if season != '':
            self.mouthBox.setItemText(0, '')
            self.mouthBox.setItemText(1, str((self.seasonInfor.index(season)-1) * 3 + 1))
            self.mouthBox.setItemText(2, str((self.seasonInfor.index(season)-1) * 3 + 2))
            self.mouthBox.setItemText(3, str((self.seasonInfor.index(season)-1) * 3 + 3))
        else:
            self.mouthBox.setItemText(0, '')
            self.mouthBox.setItemText(1, '')
            self.mouthBox.setItemText(2, '')
            self.mouthBox.setItemText(3, '')

    def initUi(self):

        self.listWidget.currentRowChanged.connect(
            self.stackedWidget.setCurrentIndex)

        self.listWidget.setFrameShape(QListWidget.NoFrame)

        self.listWidget.setVerticalScrollBarPolicy(Qt.ScrollBarAlwaysOff)
        self.listWidget.setHorizontalScrollBarPolicy(Qt.ScrollBarAlwaysOff)



        item = QListWidgetItem(str('票房份额'), self.listWidget)

        item.setSizeHint(QSize(16777215, 60))

        item.setTextAlignment(Qt.AlignCenter)

        item = QListWidgetItem(str('票房变化'), self.listWidget)
        item.setSizeHint(QSize(16777215, 60))
        item.setTextAlignment(Qt.AlignCenter)

        item = QListWidgetItem(str('票房排行'), self.listWidget)
        item.setSizeHint(QSize(16777215, 60))
        item.setTextAlignment(Qt.AlignCenter)

        item = QListWidgetItem(str('劳模演员'), self.listWidget)
        item.setSizeHint(QSize(16777215, 60))
        item.setTextAlignment(Qt.AlignCenter)

        if self.authority != 2:
            item = QListWidgetItem(str('报表生成'), self.listWidget)
            item.setSizeHint(QSize(16777215, 60))
            item.setTextAlignment(Qt.AlignCenter)

            item = QListWidgetItem(str('票房预测'), self.listWidget)
            item.setSizeHint(QSize(16777215, 60))
            item.setTextAlignment(Qt.AlignCenter)

            item = QListWidgetItem(str('电影搜索'), self.listWidget)
            item.setSizeHint(QSize(16777215, 60))
            item.setTextAlignment(Qt.AlignCenter)

    def stack1UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()
        yearlabel = QLabel("年份", self)
        yearlabel.setAlignment(Qt.AlignCenter)
        yearlabel.setMaximumSize(50,30)
        yearlabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(yearlabel)

        self.yearBox = QComboBox()
        self.yearBox.setMaximumSize(150, 50)
        # self.yearBox.setStyleSheet('color: rgb(199, 199 ,199)')
        self.yearInfor = ['2015', '2016', '2017', '2018']
        self.yearBox.addItems(self.yearInfor)
        timeLayout.addWidget(self.yearBox)

        seasonLabel = QLabel("季度", self)
        seasonLabel.setAlignment(Qt.AlignCenter)
        seasonLabel.setMaximumSize(50, 30)
        seasonLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(seasonLabel)

        timeLayout.addWidget(self.seasonBox)

        mouthLabel = QLabel("月份", self)
        mouthLabel.setAlignment(Qt.AlignCenter)
        mouthLabel.setMaximumSize(50, 30)
        mouthLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(mouthLabel)

        timeLayout.addWidget(self.mouthBox)

        typeLabel = QLabel("类型", self)
        typeLabel.setAlignment(Qt.AlignCenter)
        typeLabel.setMaximumSize(50, 30)
        typeLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(typeLabel)

        self.typeBox = QComboBox()
        self.typeBox.setMaximumSize(150, 50)
        # self.typeBox.setStyleSheet('color: rgb(199, 199 ,199)')
        self.typeInfor = ['柱状图', '饼状图']
        self.typeBox.addItems(self.typeInfor)
        timeLayout.addWidget(self.typeBox)

        if self.authority != 2:
            self.printCheck = QCheckBox("生成报表", self)
            self.printCheck.setMaximumSize(100, 30)
            self.printCheck.setMinimumSize(80, 30)
            self.printCheck.setStyleSheet('font: 10pt')
            timeLayout.addWidget(self.printCheck)

        button = QPushButton('确定', self)
        button.setMaximumSize(150,50)
        # button.setStyleSheet('color: rgb(199,199,199)')
        button.clicked.connect(self.partitionGraphBuild)
        timeLayout.addWidget(button)

        layout.addRow(timeLayout)
        self.picture = QWebEngineView(self)
        # self.picture.setAlignment(Qt.AlignCenter)
        self.picture.setMinimumSize(1500, 800)
        self.picture.setMaximumSize(1000,800)
        # picAdd = QtGui.QPixmap('./software/graph/100/2018_12_05/StyleBoxfile_2018_spring_pie.png')
        # self.picture.setStyleSheet('margin: 75px;source: ')
        layout.addRow(self.picture)

        self.stack1.setLayout(layout)

    def partitionGraphBuild(self):
        year = self.yearBox.currentText()
        season_to = {'': None, '第一季度': 'spring', '第二季度': 'summer', '第三季度': 'fall', '第四季度': 'winter'}
        season = season_to[self.seasonBox.currentText()]
        mouth = self.mouthBox.currentText()
        if mouth == '':
            mouth = None
        else:
            mouth = int(mouth)
        type_to = {'柱状图': 'histogram', '饼状图': 'pie'}
        type = type_to[self.typeBox.currentText()]
        options = {'time': [int(year), season , mouth],'visual_option': [type],  'sort_by': ['boxfile']}
        new = StyleBoxfileModel(self.id, self.date, options)
        html,add = new.process()
        html = '.\\' +html[0]
        add = '.\\' +add[0]
        if self.authority != 2:
            if (not add in self.graphList) and self.printCheck.isChecked():
                self.graphList.append(add)
        url = QUrl(QFileInfo(html).absoluteFilePath())
        picAdd = QtGui.QPixmap(add)
        # self.layout.addRow(self.picture)
        self.picture.load(url)
        del new


    def stack2UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()
        label = QLabel('请选择需要查看的年份：')
        label.setMaximumSize(250, 30)
        label.setStyleSheet('font: 12pt \"新宋体\"')
        timeLayout.addWidget(label)
        self.yearCheck = []
        for i in range(2018-2015 + 1):
            self.yearCheck.append(QCheckBox(str(2015 + i),self))
            self.yearCheck[i].setMaximumSize(100, 30)
            self.yearCheck[i].setMinimumSize(80, 30)
            self.yearCheck[i].setStyleSheet('font: 10pt')
            timeLayout.addWidget(self.yearCheck[i])

        if self.authority != 2:
            self.printCheck2 = QCheckBox("生成报表",self)
            self.printCheck2.setMaximumSize(100, 30)
            self.printCheck2.setMinimumSize(80, 30)
            self.printCheck2.setStyleSheet('font: 10pt')
            timeLayout.addWidget(self.printCheck2)

        okButton = QPushButton("确定",self)
        okButton.setMaximumSize(150,50)
        # okButton.setStyleSheet('color: rgb(199,199,199)')
        okButton.clicked.connect(self.boxChangeGraphBuild)
        timeLayout.addWidget(okButton)

        layout.addRow(timeLayout)

        self.picture2 = QWebEngineView(self)
        self.picture2.setMinimumSize(1500, 800)
        self.picture2.setMaximumSize(1000, 800)
        layout.addRow(self.picture2)


        self.stack2.setLayout(layout)

    def boxChangeGraphBuild(self):
        year = []
        for i in range(2018 - 2015 + 1):
            if self.yearCheck[i].isChecked():
                year.append(i+2015)
        options = {'years': year,'visual_option':['line']}
        new = YearBoxfileModel(self.id, self.date, options)
        html,add = new.process()
        html = '.\\' +html[0]
        add = '.\\' +add[0]
        if self.authority != 2:
            if (not add in self.graphList) and self.printCheck.isChecked() :
                self.graphList.append(add)
        url = QUrl(QFileInfo(html).absoluteFilePath())
        self.picture2.load(url)

        del new



    def stack3UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()
        yearlabel = QLabel("年份", self)
        yearlabel.setAlignment(Qt.AlignCenter)
        yearlabel.setMaximumSize(50, 30)
        yearlabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(yearlabel)

        self.yearBox3 = QComboBox()
        self.yearBox3.setMaximumSize(150, 50)
        # self.yearBox3.setStyleSheet('color: rgb(199, 199 ,199)')
        self.yearInfor3 = ['2015', '2016', '2017', '2018']
        self.yearBox3.addItems(self.yearInfor3)
        timeLayout.addWidget(self.yearBox3)

        boxLabel = QLabel("票房排名", self)
        boxLabel.setAlignment(Qt.AlignCenter)
        boxLabel.setMaximumSize(150, 30)
        boxLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(boxLabel)

        self.rankBox3 = QComboBox()
        self.rankBox3.setMaximumSize(150, 50)
        # self.rankBox3.setStyleSheet('color: rgb(199, 199 ,199)')
        self.rankInfor3 = ['1','2','3','4','5','6','7','8','9','10']
        self.rankBox3.addItems(self.rankInfor3)
        timeLayout.addWidget(self.rankBox3)

        typeLabel = QLabel("类型", self)
        typeLabel.setAlignment(Qt.AlignCenter)
        typeLabel.setMaximumSize(50, 30)
        typeLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(typeLabel)

        self.typeBox3 = QComboBox()
        self.typeBox3.setMaximumSize(150, 50)
        # self.typeBox3.setStyleSheet('color: rgb(199, 199 ,199)')
        self.typeInfor3 = ['柱状图', '云图']
        self.typeBox3.addItems(self.typeInfor3)
        timeLayout.addWidget(self.typeBox3)

        if self.authority != 2:
            self.printCheck3 = QCheckBox("生成报表", self)
            self.printCheck3.setMaximumSize(100, 30)
            self.printCheck3.setMinimumSize(80, 30)
            self.printCheck3.setStyleSheet('font: 10pt')
            timeLayout.addWidget(self.printCheck3)

        button = QPushButton('确定', self)
        button.setMaximumSize(150, 50)
        # button.setStyleSheet('color: rgb(199,199,199)')
        button.clicked.connect(self.rankingGraphBuild)
        timeLayout.addWidget(button)

        layout.addRow(timeLayout)
        self.picture3 = QWebEngineView(self)
        self.picture3.setMinimumSize(1500, 800)
        # self.picture3.setMaximumSize(1500, 1000)
        layout.addRow(self.picture3)
        self.stack3.setLayout(layout)

    def rankingGraphBuild(self):
        print(self.date)
        print(self.id)
        year = self.yearBox3.currentText()
        rank = self.rankBox3.currentText()
        type_to = {'柱状图': 'histogram', '云图': 'cloud'}
        type = type_to[self.typeBox3.currentText()]
        '''
        params: id: user_id
                logging_time: user latest logging time
                options: user's choice for visualization (type:dict)
                         {'year': a single year(int),
                          'top': int,
                          'visual_option':['cloud','histogram'],
                          'sort_by':('boxfile' or 'score')}
        '''
        options = {'year': int(year),'top': int(rank),'visual_option': [type],'sort_by': 'boxfile'}
        new = TopFilmModel(self.id, self.date, options)
        html, add = new.process()
        html = '.\\' + html[0]
        add = '.\\' + add[0]
        if self.authority != 2:
            if (not add in self.graphList) and self.printCheck.isChecked():
                self.graphList.append(add)
        url = QUrl(QFileInfo(html).absoluteFilePath())
        self.picture3.load(url)
        del new


    def stack4UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()
        yearlabel = QLabel("年份", self)
        yearlabel.setAlignment(Qt.AlignCenter)
        yearlabel.setMaximumSize(50, 30)
        yearlabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(yearlabel)

        self.yearBox4 = QComboBox()
        self.yearBox4.setMaximumSize(150, 50)
        # self.yearBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.yearInfor4 = ['2015', '2016', '2017', '2018']
        self.yearBox4.addItems(self.yearInfor4)
        timeLayout.addWidget(self.yearBox4)

        boxLabel = QLabel("排名", self)
        boxLabel.setAlignment(Qt.AlignCenter)
        boxLabel.setMaximumSize(80, 30)
        boxLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(boxLabel)

        self.rankBox4 = QComboBox()
        self.rankBox4.setMaximumSize(150, 50)
        # self.rankBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.rankInfor4 = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10']
        self.rankBox4.addItems(self.rankInfor4)
        timeLayout.addWidget(self.rankBox4)

        typeLabel = QLabel("类型", self)
        typeLabel.setAlignment(Qt.AlignCenter)
        typeLabel.setMaximumSize(50, 30)
        typeLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(typeLabel)

        self.typeBox4 = QComboBox()
        self.typeBox4.setMaximumSize(150, 50)
        # self.typeBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.typeInfor4 = ['柱状图', '云图']
        self.typeBox4.addItems(self.typeInfor4)
        timeLayout.addWidget(self.typeBox4)

        sexLabel = QLabel("性别", self)
        sexLabel.setAlignment(Qt.AlignCenter)
        sexLabel.setMaximumSize(50, 30)
        sexLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(sexLabel)

        self.sexBox4 = QComboBox()
        self.sexBox4.setMaximumSize(150, 50)
        # self.sexBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.sexInfor4 = ['男', '女']
        self.sexBox4.addItems(self.sexInfor4)
        timeLayout.addWidget(self.sexBox4)

        if self.authority != 2:
            self.printCheck4 = QCheckBox("生成报表", self)
            self.printCheck4.setMaximumSize(100, 30)
            self.printCheck4.setMinimumSize(80, 30)
            self.printCheck4.setStyleSheet('font: 10pt')
            timeLayout.addWidget(self.printCheck4)

        button = QPushButton('确定', self)
        button.setMaximumSize(150, 50)
        # button.setStyleSheet('color: rgb(199,199,199)')
        button.clicked.connect(self.actorGraphBuild)
        timeLayout.addWidget(button)

        layout.addRow(timeLayout)
        self.picture4 = QWebEngineView(self)
        self.picture4.setMinimumSize(1500, 800)
        # self.picture4.setMaximumSize(1000, 800)
        layout.addRow(self.picture4)

        self.stack4.setLayout(layout)

    def actorGraphBuild(self):
        year = self.yearBox4.currentText()
        rank = self.rankBox4.currentText()
        type_to = {'柱状图': 'histogram', '云图': 'cloud'}
        type = type_to[self.typeBox4.currentText()]
        sex_to = {'男': 'male', '女': 'female'}
        sex = sex_to[self.sexBox4.currentText()]
        options = {'year': int(year),'top': int(rank),'sex': sex,'visual_option': [type]}
        new = TopActorModel(self.id, self.date, options)
        html, add = new.process()
        html = '.\\' + html[0]
        add = '.\\' + add[0]
        if self.authority != 2:
            if (not add in self.graphList) and self.printCheck.isChecked():
                self.graphList.append(add)
        url = QUrl(QFileInfo(html).absoluteFilePath())
        self.picture4.load(url)
        del new


    def stack5UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()

        label = QLabel("打印图表",self)
        label.setAlignment(Qt.AlignCenter)
        label.setMaximumSize(150, 30)

        label.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(label)

        button = QPushButton("确定", self)
        button.setMaximumSize(150, 50)
        # button.setStyleSheet('color: rgb(199,199,199)')
        button.clicked.connect(self.pdfBuild)
        timeLayout.addWidget(button)

        layout.addRow(timeLayout)
        layout.addRow(QLabel())

        self.stack5.setLayout(layout)

    def pdfBuild(self):
        fileName, ok2 = QFileDialog.getSaveFileName(self, "文件保存", "/", "Document(*.docx)")
        print(fileName)
        print(ok2)
        doc = Document(docx=os.path.join(os.getcwd(), 'default.docx'))
        for add in self.graphList and self.graphList:
            images = add
            try:
                doc.add_picture(images, width=Inches(5))  # 添加图, 设置宽度
            except Exception:
                jpg_ima = Image.open(images)  # 打开图片
                jpg_ima.save(images)  # 保存新的图片
                doc.add_picture(images, width=Inches(5))
        if len(fileName):
            doc.save(fileName)
    def stack6UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()
        mouthlabel = QLabel("月份", self)
        mouthlabel.setAlignment(Qt.AlignCenter)
        mouthlabel.setMaximumSize(50, 30)
        mouthlabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(mouthlabel)

        self.mouthBox6 = QComboBox()
        self.mouthBox6.setMaximumSize(150, 50)
        # self.yearBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.mouthInfor4 = ['1', '2', '3', '4', '5', '6', '7', '8', '9', '10', '11', '12']
        self.mouthBox6.addItems(self.mouthInfor4)
        timeLayout.addWidget(self.mouthBox6)

        directorLabel = QLabel("导演1", self)
        directorLabel.setAlignment(Qt.AlignCenter)
        directorLabel.setMaximumSize(80, 30)
        directorLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(directorLabel)

        self.directorLine1 = QLineEdit()
        self.directorLine1.setMaximumSize(150, 50)
        # self.rankBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        timeLayout.addWidget(self.directorLine1)

        directorLabe2 = QLabel("导演2", self)
        directorLabe2.setAlignment(Qt.AlignCenter)
        directorLabe2.setMaximumSize(80, 30)
        directorLabe2.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(directorLabe2)

        self.directorLine2 = QLineEdit()
        self.directorLine2.setMaximumSize(150, 50)
        timeLayout.addWidget(self.directorLine2)

        actorLabe1 = QLabel("主演1", self)
        actorLabe1.setAlignment(Qt.AlignCenter)
        actorLabe1.setMaximumSize(80, 30)
        actorLabe1.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(actorLabe1)

        self.actorLine1 = QLineEdit()
        self.actorLine1.setMaximumSize(150, 50)
        timeLayout.addWidget(self.actorLine1)

        actorLabe2 = QLabel("主演2", self)
        actorLabe2.setAlignment(Qt.AlignCenter)
        actorLabe2.setMaximumSize(80, 30)
        actorLabe2.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(actorLabe2)

        self.actorLine2 = QLineEdit()
        self.actorLine2.setMaximumSize(150, 50)
        timeLayout.addWidget(self.actorLine2)

        typeLabel = QLabel("类型", self)
        typeLabel.setAlignment(Qt.AlignCenter)
        typeLabel.setMaximumSize(50, 30)
        typeLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(typeLabel)

        self.typeBox6 = QComboBox()
        self.typeBox6.setMaximumSize(150, 50)
        # self.typeBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.typeInfor6 = ['喜剧','奇幻','古装','动作','惊悚','犯罪','冒险','爱情','科幻','动画','剧情','悬疑',
        '战争','灾难','家庭','历史','运动','传记','武侠','歌舞','恐怖','音乐','纪录片','短片','戏曲','黑色电影','西部']
        self.typeBox6.addItems(self.typeInfor6)
        timeLayout.addWidget(self.typeBox6)


        self.printCheck6 = QCheckBox("续集？", self)
        self.printCheck6.setMaximumSize(100, 30)
        self.printCheck6.setMinimumSize(80, 30)
        self.printCheck6.setStyleSheet('font: 10pt')
        timeLayout.addWidget(self.printCheck6)

        button = QPushButton('确定', self)
        button.setMaximumSize(150, 50)
        # button.setStyleSheet('color: rgb(199,199,199)')
        button.clicked.connect(self.boxPre)
        timeLayout.addWidget(button)

        layout.addRow(timeLayout)
        self.picture6 = QLabel()
        self.picture6.setAlignment(Qt.AlignCenter)
        self.picture6.setMinimumSize(300, 600)
        # self.picture6.setMaximumSize(1200, 800)
        self.picture6.setStyleSheet('margin: 75px;source: ')
        layout.addRow(self.picture6)

        self.stack6.setLayout(layout)

    def boxPre(self):
        month = self.mouthBox6.currentText()
        mouth = int(month)
        direct = []
        actor = []
        if len(self.directorLine1.text()) > 0:
            direct.append(self.directorLine1.text())
        if len(self.directorLine2.text()) > 0:
            direct.append(self.directorLine2.text())
        if len(self.actorLine1.text()) > 0:
            actor.append(self.actorLine1.text())
        if len(self.actorLine2.text()) > 0:
            actor.append(self.actorLine2.text())
        type = self.typeBox6.currentText()
        isSequel = self.printCheck6.isChecked()
        if actor == [] or direct == []:
            QMessageBox.information(self, '提示信息', '导演和演员至少要各输入一个！')
        else:
            options = {'month': mouth ,'directors': direct,'actors': actor,'genres': [type],'isSequel': isSequel}
            new = PredictModel(options)
            prediction = round(new.predict(),2)
            QMessageBox.information(self, '预测结果', str(prediction) + '百万')
            self.picture6.setText(str(prediction) + '百万')
            self.picture6.setStyleSheet('font: 50pt \"黑体\"; color: rgb(255, 200, 1);border: 2px solid rgb(255, 200, 1); ')
            del new

    def stack7UI(self):
        layout = QFormLayout()
        timeLayout = QHBoxLayout()
        typeLabel = QLabel("搜索类型", self)
        typeLabel.setAlignment(Qt.AlignCenter)
        typeLabel.setMaximumSize(100, 30)
        typeLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(typeLabel)

        self.typeBox7 = QComboBox()
        self.typeBox7.setMaximumSize(150, 50)
        # self.yearBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.typeInfor7 = ['电影名字', '导演名字', '演员名字', '电影类型']
        self.typeBox7.addItems(self.typeInfor7)
        timeLayout.addWidget(self.typeBox7)

        Label = QLabel("关键字", self)
        Label.setAlignment(Qt.AlignCenter)
        Label.setMaximumSize(100, 30)
        Label.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(Label)

        self.searchLine = QLineEdit()
        self.searchLine.setMaximumSize(150, 50)
        timeLayout.addWidget(self.searchLine)

        dataLabel = QLabel("搜索依据", self)
        dataLabel.setAlignment(Qt.AlignCenter)
        dataLabel.setMaximumSize(100, 30)
        dataLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(dataLabel)

        self.dataBox7 = QComboBox()
        self.dataBox7.setMaximumSize(150, 50)
        # self.yearBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.dataInfor7 = ['电影评分', '电影票房']
        self.dataBox7.addItems(self.dataInfor7)
        timeLayout.addWidget(self.dataBox7)

        numberLabel = QLabel("排序方式", self)
        numberLabel.setAlignment(Qt.AlignCenter)
        numberLabel.setMaximumSize(100, 30)
        numberLabel.setStyleSheet('font: 15pt \"新宋体\"')
        timeLayout.addWidget(numberLabel)

        self.numberBox7 = QComboBox()
        self.numberBox7.setMaximumSize(150, 50)
        # self.yearBox4.setStyleSheet('color: rgb(199, 199 ,199)')
        self.numberinfo7 = ['降序', '升序']
        self.numberBox7.addItems(self.numberinfo7)
        timeLayout.addWidget(self.numberBox7)

        button = QPushButton('确定', self)
        button.setMaximumSize(150, 50)
        # button.setStyleSheet('color: rgb(199,199,199)')
        button.clicked.connect(self.search)
        timeLayout.addWidget(button)

        layout.addRow(timeLayout)

        self.tableWidget = QTableWidget()
        self.tableWidget.setObjectName("tableWidget")
        self.tableWidget.setMinimumSize(150, 700)
        # self.tableWidget.setMaximumSize(1050,800)
        self.tableWidget.setColumnCount(8)
        self.tableWidget.setRowCount(0)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(0, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(1, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(2, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(3, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(4, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(5, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(6, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(7, item)
        item = QTableWidgetItem()
        self.tableWidget.setHorizontalHeaderItem(8, item)
        self.tableWidget.clear()
        # 重新设置表头
        self.tableWidget.setHorizontalHeaderLabels(
            ['电影名字', '上映年份', '上映月份', '票房', '导演', '演员', '类型', '评分'])
        self.tableWidget.horizontalHeader().setSectionResizeMode(QHeaderView.Stretch)
        self.tableWidget.horizontalHeader().setSectionResizeMode(5, QHeaderView.ResizeToContents)
        self.tableWidget.horizontalHeader().setSectionResizeMode(7, QHeaderView.ResizeToContents)
        layout.addRow(self.tableWidget)

        self.stack7.setLayout(layout)

    def search(self):
        type = self.typeBox7.currentText()
        type_to =  {'电影名字': 'name', '导演名字': 'director', '演员名字': 'actor', '电影类型': 'genre'}
        type = type_to[type]
        content = self.searchLine.text()
        if len(content) == 0:
            QMessageBox.information(self, '提示信息', '你没有输入哦！')
        sortby_to = {'电影评分': 'mark', '电影票房': 'boxfile'}
        sortby = sortby_to[self.dataBox7.currentText()]
        if self.numberBox7.currentText() == '降序':
            reverse = True
        else:
            reverse = False
        # length = int(self.numberBox7.currentText())
        if len(content) > 0:
            new = sift()
            df = new.find(type,content,sortby,reverse)
            self.tableWidget.setRowCount(len(df))
            for row in range(len(df)):
                self.tableWidget.setItem(row, 0, QTableWidgetItem(df[row]['name']))
                self.tableWidget.setItem(
                    row, 1, QTableWidgetItem(str(df[row]['year'])))
                self.tableWidget.setItem(
                    row, 2, QTableWidgetItem(str(df[row]['month'])))
                self.tableWidget.setItem(
                    row, 3, QTableWidgetItem(str(df[row]['boxfile'])))
                self.tableWidget.setItem(
                    row, 4, QTableWidgetItem(",".join(df[row]['director'])))
                self.tableWidget.setItem(
                    row, 5, QTableWidgetItem(",".join(df[row]['actor'])))
                self.tableWidget.setItem(
                    int(row), 6, QTableWidgetItem(",".join(df[row]['genre'])))
                self.tableWidget.setItem(int(row), 7, QTableWidgetItem(str(df[row]['score'])))
            self.tableWidget.resizeColumnsToContents()
            self.tableWidget.resizeRowsToContents()
            del new
            del df




import source_rc




if __name__ == '__main__':
    import sys
    from PyQt5.QtWidgets import QApplication
    app = QApplication(sys.argv)
    w = LeftTabWidget(1, '2018_12_18', 1)

    w.show()
    sys.exit(app.exec_())
